import os
import logging
import pdfplumber
from docx import Document

logging.basicConfig(level=logging.INFO)


def extract_text_pdf(pdf_path):
    """
    Extract text from a PDF file.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Extracted text.
    """
    text = ""

    with pdfplumber.open(pdf_path) as pdf:
        logging.info(f"PDF pages detected: {len(pdf.pages)}")

        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    return text.strip()


def extract_text_doc(doc_path):
    """
    Extract text from a DOCX file.

    Args:
        doc_path (str): Path to the DOCX file.

    Returns:
        str: Extracted text.
    """
    text = ""

    doc = Document(doc_path)
    logging.info(f"DOCX paragraphs detected: {len(doc.paragraphs)}")

    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"

    return text.strip()


def extract_resume_text(file_path):
    """
    Detect file type and extract text accordingly.

    Args:
        file_path (str): Path to PDF or DOCX file.

    Returns:
        str: Extracted text.

    Raises:
        ValueError: If file format is unsupported or text is empty.
    """

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    name, ext = os.path.splitext(file_path)
    ext = ext.lower()

    logging.info(f"Detected file type: {ext}")

    if ext == ".pdf":
        text = extract_text_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_doc(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX allowed.")

    if not text.strip():
        raise ValueError("No text could be extracted from the file.")

    return text
